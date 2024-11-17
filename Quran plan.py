from datetime import date, timedelta
from docx import Document
from docx.shared import Pt, Inches

# Generate dates for the 2-month plan
start_date = date.today()
end_date = start_date + timedelta(days=60)
current_date = start_date

# Initialize Word Document
doc = Document()
doc.add_heading("Memorization Tracker for 30th Juz", level=1)
doc.add_paragraph(f"Start Date: {start_date} | End Date: {end_date}\n")

# Memorization and Revision Plan Details
plan_details = [
    "Morning Memorization: Focus on new Ayahs.",
    "Evening Revision: Revise previously memorized sections.",
    "Weekly Review: Consolidate all memorized Surahs and address any challenges.",
    "Track accuracy, challenges, and improvements daily.",
]

doc.add_heading("Plan Overview", level=2)
for detail in plan_details:
    doc.add_paragraph(detail, style="List Bullet")

# Create a daily tracker table
doc.add_heading("Daily Memorization Tracker", level=2)
table = doc.add_table(rows=1, cols=6)
table.style = "Table Grid"

# Add headers
headers = ["Day", "Date", "Surah & Ayah", "Morning Memorization (✓)", "Evening Revision (✓)", "Notes/Challenges"]
for idx, header in enumerate(headers):
    table.cell(0, idx).text = header

# Fill the table for 60 days
surahs = [
    ("Al-Nas", 6), ("Al-Falaq", 5), ("Al-Ikhlas", 4), ("Al-Masad", 5), ("Al-Nasr", 3), ("Al-Kafirun", 6),
    ("Al-Kawthar", 3), ("Al-Ma'un", 7), ("Quraish", 4), ("Al-Fil", 5), ("Al-Humazah", 9), ("Al-Asr", 3),
    ("At-Takathur", 8), ("Al-Qari'ah", 11), ("Al-Adiyat", 11), ("Az-Zalzalah", 8), ("Al-Bayyinah", 8),
    ("Al-Qadr", 5), ("Al-Tin", 8), ("Al-Alaq", 19), ("Al-Sharh", 8), ("Al-Duha", 11), ("Al-Lail", 21),
    ("Ash-Shams", 15), ("Al-Balad", 20), ("Al-Fajr", 30), ("Al-Ghashiyah", 26), ("Al-A'la", 19),
    ("At-Tariq", 17), ("Al-Buruj", 22), ("Al-Inshiqaq", 25), ("Al-Mutaffifin", 36), ("Al-Infitar", 19),
    ("An-Nazi'at", 46), ("An-Naba", 40)
]

day = 1
surah_idx = 0
ayah_start = 1

while current_date <= end_date:
    if surah_idx >= len(surahs):
        break

    surah, total_ayahs = surahs[surah_idx]
    ayah_end = min(ayah_start + 3 - 1, total_ayahs)  # Memorize up to 3 Ayahs/day

    # Add row for the day
    row = table.add_row().cells
    row[0].text = f"Day {day}"
    row[1].text = current_date.strftime("%Y-%m-%d")
    row[2].text = f"{surah} - Ayahs {ayah_start}-{ayah_end}"
    row[3].text = ""  # Placeholder for Morning Memorization
    row[4].text = ""  # Placeholder for Evening Revision
    row[5].text = ""  # Placeholder for Notes/Challenges

    # Update for next iteration
    ayah_start = ayah_end + 1
    if ayah_start > total_ayahs:
        ayah_start = 1
        surah_idx += 1

    current_date += timedelta(days=1)
    day += 1

# Add weekly summary section
doc.add_heading("Weekly Progress Summary", level=2)
for week in range(1, 9):
    doc.add_heading(f"Week {week}", level=3)
    summary_table = doc.add_table(rows=2, cols=6)
    summary_table.style = "Table Grid"

    # Headers for weekly progress
    summary_headers = [
        "Surahs Covered", "Total Ayahs Memorized", "Reviewed Surahs",
        "Challenges Faced", "Solutions/Improvements", "Motivational Notes"
    ]
    for idx, header in enumerate(summary_headers):
        summary_table.cell(0, idx).text = header

    for cell in summary_table.rows[1].cells:
        cell.text = ""

# Save document
file_path = "/mnt/data/Memorization_Tracker_2_Months.docx"
doc.save(file_path)
file_path

